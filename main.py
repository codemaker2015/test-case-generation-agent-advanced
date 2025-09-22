import io
import os
import pypdf 
import json
import datetime
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
import pandas as pd
from io import BytesIO
from agent import initialize_app
from langchain_groq.chat_models import ChatGroq
from docx import Document

# Page configuration
st.set_page_config(
    page_title=" Test Case Generation Agent",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
<style>
    .main-header {
        padding: 1rem 0;
        border-bottom: 2px solid #e0e0e0;
        margin-bottom: 2rem;
    }
    .metric-card {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 10px;
        border: 1px solid #dee2e6;
        margin-bottom: 1rem;
    }
    .stage-indicator {
        background-color: #e3f2fd;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        border-left: 4px solid #2196f3;
        margin: 0.5rem 0;
    }
    .test-case-section {
        background-color: #f1f8e9;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        border-left: 4px solid #4caf50;
    }
    .risk-high {
        color: #d32f2f;
        font-weight: bold;
    }
    .risk-medium {
        color: #f57c00;
        font-weight: bold;
    }
    .risk-low {
        color: #388e3c;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)

st.title("🔬  Test Case Generation Agent")
st.markdown("*AI-Powered Comprehensive Testing Solution with Risk Analysis & Coverage Optimization*")

# Initialize session state for messages and  features
if "messages" not in st.session_state:
    st.session_state.messages = []
if "test_metrics" not in st.session_state:
    st.session_state.test_metrics = {}
if "current_session_data" not in st.session_state:
    st.session_state.current_session_data = {}

#  Sidebar Configuration
with st.sidebar:
    st.header("🛠️ Configuration")
    
    # File upload section
    st.subheader("📄 Requirements Document")
    uploaded_file = st.file_uploader(
        "Upload Requirements Document", 
        type=["txt", "pdf", "docx"],
        help="Upload your requirements document for analysis"
    )
    if "uploaded_file" not in st.session_state or uploaded_file is not None:
        st.session_state.uploaded_file = uploaded_file

    # Model selection with  options
    st.subheader("🤖 AI Model Configuration")
    model_options = [
        "llama-3.1-8b-instant",
        "llama-3.3-70b-versatile",
        "llama3-70b-8192",
        "llama3-8b-8192",
        "mixtral-8x7b-32768",
        "gemma2-9b-it"
    ]
    
    if "selected_model" not in st.session_state:
        st.session_state.selected_model = "llama-3.1-8b-instant"
            
    selected_model = st.selectbox(
        "Select AI Model", 
        model_options, 
        key="selected_model", 
        index=model_options.index(st.session_state.selected_model),
        help="Choose the AI model for test case generation"
    )
    
    # Update model when changed
    if selected_model != st.session_state.selected_model:
        st.session_state.selected_model = selected_model
        st.session_state.llm = ChatGroq(model=selected_model, temperature=0.0)
    
    if "llm" not in st.session_state:
        st.session_state.llm = ChatGroq(model=selected_model, temperature=0.0)

    #  Test Generation Settings
    st.subheader("⚙️ Test Generation Settings")
    
    # Test format selection
    test_formats = ["Auto-detect", "Gherkin (BDD)", "Selenium (UI)", "API Testing", "Performance Testing", "Security Testing"]
    test_format = st.selectbox(
        "Preferred Test Format", 
        test_formats, 
        index=0,
        help="Choose the default test format or let AI auto-detect"
    )
    st.session_state.test_format = test_format
    
    # Advanced settings in expandable sections
    with st.expander("🎯 Coverage & Quality Settings"):
        include_edge_cases = st.checkbox("Include Edge Cases", value=True)
        include_negative_tests = st.checkbox("Include Negative Test Scenarios", value=True)
        include_performance_tests = st.checkbox("Include Performance Considerations", value=False)
        include_security_tests = st.checkbox("Include Security Test Cases", value=False)
        
        st.session_state.include_edge_cases = include_edge_cases
        st.session_state.include_negative_tests = include_negative_tests
        st.session_state.include_performance_tests = include_performance_tests
        st.session_state.include_security_tests = include_security_tests
    
    with st.expander("📊 Analysis & Reporting"):
        enable_risk_analysis = st.checkbox("Enable Risk Assessment", value=True)
        enable_coverage_analysis = st.checkbox("Enable Coverage Analysis", value=True)
        enable_defect_prediction = st.checkbox("Enable Defect Prediction", value=True)
        generate_execution_plan = st.checkbox("Generate Execution Plan", value=True)
        
        st.session_state.enable_risk_analysis = enable_risk_analysis
        st.session_state.enable_coverage_analysis = enable_coverage_analysis
        st.session_state.enable_defect_prediction = enable_defect_prediction
        st.session_state.generate_execution_plan = generate_execution_plan
    
    with st.expander("🔧 Advanced Options"):
        enhancement_level = st.slider("Test Case Detail Level", min_value=1, max_value=5, value=3)
        use_industry_standards = st.checkbox("Apply Industry Best Practices", value=True)
        enable_traceability = st.checkbox("Generate Requirements Traceability", value=True)
        
        st.session_state.enhancement_level = enhancement_level
        st.session_state.use_industry_standards = use_industry_standards
        st.session_state.enable_traceability = enable_traceability
    
    # Session management
    st.subheader("🔄 Session Management")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🗑️ Clear Chat", use_container_width=True):
            st.session_state.messages = []
            st.rerun()
    with col2:
        if st.button("📊 Show Metrics", use_container_width=True):
            st.session_state.show_metrics = True

# Initialize the  app
app = initialize_app(model_name=st.session_state.selected_model)

# Main content area with tabs
tab1, tab2, tab3, tab4 = st.tabs(["💬 Chat Interface", "📊 Analytics Dashboard", "📋 Test Reports", "ℹ️ Help & Examples"])

with tab1:
    #  help section
    with st.expander("🎯 How to use this  Test Generation Agent", expanded=False):
        st.markdown("""
        ### 🚀  Test Case Generation Agent
        
        This advanced AI agent provides comprehensive test case generation with:
        
        **📋 Core Features:**
        - **Multi-format Test Generation**: Gherkin, Selenium, API, Performance, Security
        - **Risk Assessment**: Identifies high-risk areas requiring thorough testing
        - **Coverage Analysis**: Ensures comprehensive test coverage
        - **Defect Prediction**: AI-powered predictions of potential defect areas
        - **Test Prioritization**: Smart prioritization based on risk and business impact
        - **Requirements Traceability**: Links test cases to requirements
        
        **🔧 How to Use:**
        1. **Upload Requirements**: Use the sidebar to upload your requirements document
        2. **Configure Settings**: Adjust test generation preferences
        3. **Ask Natural Questions**: Examples below
        4. **Review Results**: Get comprehensive test reports with analytics
        
        **💡 Example Requests:**
        - *"Generate comprehensive test cases for the user authentication system"*
        - *"Create API test cases for the payment processing module with security focus"*
        - *"Generate Gherkin scenarios for e-commerce checkout with edge cases"*
        - *"Analyze requirements and create a complete test plan with risk assessment"*
        - *"Generate performance test cases for the search functionality"*
        """)

    # Display chat messages with  formatting
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            if message["role"] == "assistant" and "test_metrics" in message:
                # Display  response with metrics
                st.markdown(message["content"])
                
                # Show metrics if available
                if message.get("test_metrics"):
                    with st.expander("📊 Generation Metrics"):
                        metrics = message["test_metrics"]
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("Test Cases", metrics.get("total_test_cases", 0))
                        with col2:
                            st.metric("Format", metrics.get("format_used", "N/A"))
                        with col3:
                            st.metric("Validation", "✅" if metrics.get("validation_performed") == "Yes" else "❌")
                        with col4:
                            st.metric("Coverage", "✅" if metrics.get("coverage_analyzed") == "Yes" else "❌")
            else:
                st.markdown(message["content"])

    # Get requirements document content
    requirements_docs_content = ""
    if "uploaded_file" in st.session_state and st.session_state.uploaded_file is not None:
        try:
            if st.session_state.uploaded_file.type == "text/plain":
                requirements_docs_content = st.session_state.uploaded_file.getvalue().decode("utf-8")
            elif st.session_state.uploaded_file.type == "application/pdf":
                pdf_reader = pypdf.PdfReader(io.BytesIO(st.session_state.uploaded_file.getvalue()))
                for page in pdf_reader.pages:
                    requirements_docs_content += page.extract_text()
        except Exception as e:
            st.error(f"Error reading uploaded file: {e}")
    elif os.path.exists("./input.txt"):
        try:
            with open("./input.txt", "r", encoding='utf-8') as f:
                requirements_docs_content = f.read()
        except Exception as e:
            st.error(f"Error reading default file: {e}")

    #  chat input
    user_request = st.chat_input("🎯 Enter your testing request (e.g., 'Generate comprehensive test cases with risk analysis'):")

    if user_request:
        # Add user request to chat history
        st.session_state.messages.append({"role": "user", "content": user_request})
        with st.chat_message("user"):
            st.markdown(user_request)

        # Process with  AI workflow
        with st.chat_message("assistant"):
            with st.spinner("🚀 Processing with  AI Pipeline..."):
                
                # Prepare  inputs
                format_mapping = {
                    "Auto-detect": "auto",
                    "Gherkin (BDD)": "gherkin",
                    "Selenium (UI)": "selenium", 
                    "API Testing": "api",
                    "Performance Testing": "performance",
                    "Security Testing": "security"
                }
                selected_format = format_mapping.get(st.session_state.test_format, "auto")

                inputs = {
                    "user_request": user_request,
                    "requirements_docs_content": requirements_docs_content,
                    "testcases_format": selected_format
                }
                
                #  progress tracking
                progress_placeholder = st.empty()
                response_placeholder = st.empty()
                total_answer = ""
                
                # Define comprehensive stages
                stages = [
                    "📋 Analyzing Requirements",
                    "⚠️ Assessing Risks", 
                    "🔍 Researching Best Practices",
                    "🧪 Generating Test Cases",
                    "✅ Validating Test Cases",
                    "📊 Analyzing Coverage",
                    "🔮 Predicting Defects",
                    "📝 Creating Execution Plan",
                    "📑 Generating Final Report"
                ]
                
                current_stage = 0
                stage_progress = {}
                
                # Create progress display
                progress_container = st.container()
                with progress_container:
                    progress_cols = st.columns(len(stages))
                
                # Stream  results
                for output in app.stream(inputs):
                    for node_name, state in output.items():
                        # Update stage based on node
                        stage_mapping = {
                            "_summary_node": 0,
                            "risk_analysis_node": 1,
                            "_best_practices_node": 2,
                            "gherkin_node": 3, "selenium_node": 3, "api_node": 3, 
                            "performance_node": 3, "security_node": 3,
                            "validation_node": 4,
                            "coverage_analysis_node": 5,
                            "defect_prediction_node": 6,
                            "prioritization_node": 7,
                            "final_report_node": 8
                        }
                        
                        if node_name in stage_mapping:
                            current_stage = stage_mapping[node_name]
                            stage_progress[current_stage] = True
                            
                            # Update progress display
                            with progress_placeholder.container():
                                cols = st.columns(len(stages))
                                for i, stage in enumerate(stages):
                                    with cols[i]:
                                        if i in stage_progress:
                                            st.success(f"✅ {stage.split(' ', 1)[1]}")
                                        elif i == current_stage:
                                            st.info(f"⏳ {stage.split(' ', 1)[1]}")
                                        else:
                                            st.write(f":gray[⭕ {stage.split(' ', 1)[1]}]")
                        
                        # Update response content
                        if 'answer' in state or 'final_report' in state:
                            total_answer = state.get('final_report') or state.get('answer', '')
                            response_placeholder.markdown(total_answer)
                            
                            # Store session data for analytics
                            if 'test_metrics' in state:
                                st.session_state.current_session_data = {
                                    'test_metrics': state['test_metrics'],
                                    'validation_results': state.get('validation_results', ''),
                                    'coverage_analysis': state.get('coverage_analysis', ''),
                                    'risk_assessment': state.get('risk_assessment', ''),
                                    'defect_prediction': state.get('defect_prediction', ''),
                                    'timestamp': datetime.datetime.now().isoformat()
                                }

                # Clear progress when done
                progress_placeholder.empty()
                
                # Add  response to chat history
                response_data = {
                    "role": "assistant", 
                    "content": total_answer,
                    "test_metrics": st.session_state.current_session_data.get('test_metrics', {})
                }
                st.session_state.messages.append(response_data)

with tab2:
    st.header("📊 Analytics Dashboard")
    
    if st.session_state.current_session_data:
        data = st.session_state.current_session_data
        metrics = data.get('test_metrics', {})
        
        # Key Metrics Row
        st.subheader("🎯 Key Metrics")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric(
                "Total Test Cases", 
                metrics.get('total_test_cases', 0),
                help="Total number of generated test cases"
            )
        with col2:
            st.metric(
                "Test Format", 
                metrics.get('format_used', 'N/A').upper(),
                help="Format used for test case generation"
            )
        with col3:
            validation_status = "✅ Passed" if metrics.get('validation_performed') == 'Yes' else "❌ Skipped"
            st.metric(
                "Validation", 
                validation_status,
                help="Whether test cases were validated"
            )
        with col4:
            coverage_status = "✅ Analyzed" if metrics.get('coverage_analyzed') == 'Yes' else "❌ Skipped"
            st.metric(
                "Coverage Analysis", 
                coverage_status,
                help="Whether coverage analysis was performed"
            )
        
        # Visual Analytics
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📈 Test Generation Timeline")
            timeline_data = {
                'Stage': ['Requirements', 'Risk Analysis', 'Best Practices', 'Generation', 'Validation', 'Analysis'],
                'Completed': [1, 1, 1, 1, 1 if metrics.get('validation_performed') == 'Yes' else 0, 1 if metrics.get('coverage_analyzed') == 'Yes' else 0]
            }
            df_timeline = pd.DataFrame(timeline_data)
            fig_timeline = px.bar(df_timeline, x='Stage', y='Completed', 
                                title="Pipeline Completion Status",
                                color='Completed', 
                                color_continuous_scale='RdYlGn')
            st.plotly_chart(fig_timeline, use_container_width=True)
        
        with col2:
            st.subheader("🎭 Test Case Distribution")
            if metrics.get('total_test_cases', 0) > 0:
                # Mock data for demonstration - in real implementation, parse actual test cases
                test_types = ['Functional', 'Edge Cases', 'Negative', 'Integration', 'Security']
                test_counts = [
                    int(metrics.get('total_test_cases', 0) * 0.4),
                    int(metrics.get('total_test_cases', 0) * 0.25),
                    int(metrics.get('total_test_cases', 0) * 0.2),
                    int(metrics.get('total_test_cases', 0) * 0.1),
                    int(metrics.get('total_test_cases', 0) * 0.05)
                ]
                
                fig_pie = px.pie(values=test_counts, names=test_types, 
                               title="Test Case Type Distribution")
                st.plotly_chart(fig_pie, use_container_width=True)
            else:
                st.info("No test case data available for visualization")
        
        # Detailed Analysis Sections
        if data.get('risk_assessment'):
            with st.expander("⚠️ Risk Assessment Details"):
                st.markdown(data['risk_assessment'])
        
        if data.get('coverage_analysis'):
            with st.expander("📊 Coverage Analysis Details"):
                st.markdown(data['coverage_analysis'])
        
        if data.get('defect_prediction'):
            with st.expander("🔮 Defect Prediction Details"):
                st.markdown(data['defect_prediction'])
    
    else:
        st.info("Generate test cases in the Chat Interface to see analytics here.")
        st.markdown("### 📈 Sample Analytics Preview")
        st.markdown("Once you generate test cases, you'll see:")
        st.markdown("- **Test Metrics**: Count, format, validation status")
        st.markdown("- **Visual Charts**: Timeline, distribution, coverage")
        st.markdown("- **Detailed Analysis**: Risk assessment, predictions, recommendations")

with tab3:
    st.header("📋 Test Reports")
    
    if st.session_state.current_session_data:
        data = st.session_state.current_session_data
        
        # Report Header
        st.subheader("📑 Latest Test Generation Report")
        st.markdown(f"**Generated:** {data.get('timestamp', 'Unknown')}")
        
        # Download Report
        if st.button("📥 Download Full Report"):
            report_content = ""
            for msg in st.session_state.messages:
                if msg["role"] == "assistant":
                    report_content = msg["content"]
                    break
            
            if report_content:
                doc = Document()
                doc.add_heading("Test Case Generation Report", level=1)
                doc.add_paragraph(report_content)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="📄 Download as Docx File",
                    data=buffer,
                    file_name=f"test_report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.download_button(
                    label="📄 Download as Text File",
                    data=report_content,
                    file_name=f"test_report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain"
                )
        
        # Report Summary
        metrics = data.get('test_metrics', {})
        
        # Executive Summary
        st.markdown("### 📊 Executive Summary")
        summary_cols = st.columns(3)
        with summary_cols[0]:
            st.markdown(f"**Test Cases Generated:** {metrics.get('total_test_cases', 0)}")
            st.markdown(f"**Format Used:** {metrics.get('format_used', 'N/A').upper()}")
        with summary_cols[1]:
            st.markdown(f"**Validation:** {metrics.get('validation_performed', 'No')}")
            st.markdown(f"**Coverage Analysis:** {metrics.get('coverage_analyzed', 'No')}")
        with summary_cols[2]:
            st.markdown(f"**Generation Time:** {data.get('timestamp', 'Unknown')}")
            st.markdown(f"**Requirements:** {metrics.get('requirements_analyzed', 'No')}")
        
        # Report Sections
        sections = [
            ("🎯 Risk Assessment", data.get('risk_assessment')),
            ("✅ Validation Results", data.get('validation_results')),
            ("📊 Coverage Analysis", data.get('coverage_analysis')),
            ("🔮 Defect Predictions", data.get('defect_prediction'))
        ]
        
        for title, content in sections:
            if content:
                with st.expander(title):
                    st.markdown(content)
    
    else:
        st.info("No test reports available. Generate test cases first to see reports here.")
        st.markdown("### 📋 Available Report Features")
        st.markdown("- **Executive Summary**: Key metrics and overview")
        st.markdown("- **Risk Assessment**: Identified risks and mitigation strategies")
        st.markdown("- **Validation Results**: Test case quality analysis")
        st.markdown("- **Coverage Analysis**: Requirements coverage mapping")
        st.markdown("- **Defect Predictions**: AI-powered defect prone area analysis")
        st.markdown("- **Download Options**: Export reports in multiple formats")

with tab4:
    st.header("ℹ️ Help & Examples")
    
    # Feature Overview
    st.markdown("## 🌟  Features")
    
    features = [
        ("🧪 Multi-Format Test Generation", "Supports Gherkin, Selenium, API, Performance, and Security test formats"),
        ("⚠️ Risk Assessment", "AI-powered analysis to identify high-risk areas requiring thorough testing"),
        ("📊 Coverage Analysis", "Comprehensive analysis ensuring all requirements are covered"),
        ("🔮 Defect Prediction", "Machine learning insights to predict potential defect-prone areas"),
        ("📋 Requirements Traceability", "Links each test case back to specific requirements"),
        ("🎯 Test Prioritization", "Smart prioritization based on risk and business impact"),
        ("📈 Analytics Dashboard", "Visual insights and metrics about your test generation"),
        ("📑 Comprehensive Reports", "Detailed reports with actionable insights")
    ]
    
    for title, description in features:
        st.markdown(f"**{title}**: {description}")
    
    st.markdown("## 💡 Example Requests")
    
    examples = [
        {
            "title": "🔐 Authentication System Testing",
            "request": "Generate comprehensive test cases for the user authentication system including edge cases and security considerations",
            "description": "Creates BDD scenarios, edge cases, security tests, and risk analysis for authentication"
        },
        {
            "title": "💳 Payment Processing",
            "request": "Create API test cases for payment processing with performance and security focus",
            "description": "Generates API tests, performance scenarios, security validations, and error handling"
        },
        {
            "title": "🛒 E-commerce Checkout",
            "request": "Generate Selenium test cases for checkout process with comprehensive coverage analysis",
            "description": "Creates UI automation tests with coverage mapping and validation"
        },
        {
            "title": "📊 Complete Test Strategy",
            "request": "Analyze requirements and create a complete test plan with risk assessment and execution strategy",
            "description": "Full test planning with risk analysis, prioritization, and execution roadmap"
        }
    ]
    
    for example in examples:
        with st.expander(f"📝 {example['title']}"):
            st.code(example['request'], language="text")
            st.markdown(f"**What this generates:** {example['description']}")
    
    st.markdown("## 🔧 Advanced Configuration")
    
    st.markdown("""
    ### ⚙️ Settings Explained
    
    **Test Format Options:**
    - **Auto-detect**: AI determines the best format based on your request
    - **Gherkin (BDD)**: Behavior-driven development scenarios
    - **Selenium (UI)**: Web automation test scripts
    - **API Testing**: REST/GraphQL endpoint testing
    - **Performance**: Load and stress testing scenarios
    - **Security**: Security-focused test cases
    
    **Coverage & Quality:**
    - **Edge Cases**: Boundary conditions and unusual scenarios
    - **Negative Tests**: Error conditions and invalid inputs
    - **Performance**: Response time and load considerations
    - **Security**: Authentication, authorization, data protection
    
    **Analysis & Reporting:**
    - **Risk Assessment**: Identifies high-risk areas
    - **Coverage Analysis**: Maps tests to requirements
    - **Defect Prediction**: AI insights on potential issues
    - **Execution Plan**: Prioritized test execution strategy
    """)
    
    st.markdown("## 🎯 Best Practices")
    
    st.markdown("""
    ### 📋 For Best Results:
    
    1. **Upload Clear Requirements**: Provide detailed, well-structured requirements documents
    2. **Be Specific in Requests**: Mention specific features, user journeys, or testing focus areas
    3. **Configure Settings**: Adjust settings based on your testing needs and project requirements
    4. **Review Analytics**: Use the dashboard to understand coverage and identify gaps
    5. **Download Reports**: Save comprehensive reports for documentation and team sharing
    
    ### 🚀 Getting Started Tips:
    - Start with simple requests to understand the system
    - Experiment with different test formats
    - Use the risk assessment to prioritize testing efforts
    - Leverage the execution plan for test scheduling
    """)